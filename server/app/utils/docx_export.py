from __future__ import annotations

from io import BytesIO

from docx import Document

from ..models import DocumentSection, Project


def build_docx(project: Project, sections: list[DocumentSection]) -> BytesIO:
    document = Document()
    document.add_heading(project.title, 0)
    document.add_paragraph(project.topic)
    document.add_page_break()

    for section in sorted(sections, key=lambda s: s.position):
        document.add_heading(section.title, level=1)
        document.add_paragraph(section.content or "Content pending.")
        document.add_page_break()

    buffer = BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return buffer

